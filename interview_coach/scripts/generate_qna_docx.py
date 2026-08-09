#!/usr/bin/env python3
"""
Generate a professional Word (.docx) Q&A bank from the Neon database.
Organises all questions by Topic → Subtopic → Difficulty, with:
  - Cover page
  - Table of contents
  - Numbered Q&As with full answers, keywords, and difficulty badge
  - Source attribution appendix

Run after scrape_and_seed.py has been executed.
"""

import sys, os, re
sys.path.append(os.path.join(os.path.dirname(__file__), ".."))

from collections import defaultdict
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from database.connection import SessionLocal
from database.models import Questions


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

DIFF_BADGE = {
    "easy":   "🟢 Easy",
    "medium": "🟡 Medium",
    "hard":   "🟠 Hard",
    "expert": "🔴 Expert",
}

DIFF_COLOR = {
    "easy":   RGBColor(0x2E, 0x86, 0x0E),   # green
    "medium": RGBColor(0xB8, 0x86, 0x00),   # amber
    "hard":   RGBColor(0xC4, 0x55, 0x00),   # orange
    "expert": RGBColor(0xC0, 0x00, 0x00),   # red
}


def _set_heading_color(paragraph, rgb: RGBColor):
    for run in paragraph.runs:
        run.font.color.rgb = rgb


def _add_horizontal_rule(doc: Document):
    """Insert a thin horizontal line (border paragraph)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pb.append(bottom)
    pr.append(pb)


def _clean_md_for_docx(text: str) -> str:
    """Strip markdown formatting for plain docx paragraphs."""
    # Remove image tags
    text = re.sub(r'!\[.*?\]\(.*?\)', '[See diagram in source reference]', text)
    # Inline links → keep text
    text = re.sub(r'\[([^\]]+)\]\(https?://[^)]+\)', r'\1', text)
    # Bold/italic
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',     r'\1', text)
    text = re.sub(r'__(.+?)__',     r'\1', text)
    # Inline code → keep as-is
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text.strip()


def _add_answer_block(doc: Document, raw_answer: str):
    """
    Parse the answer text and add it to the document with proper formatting.
    Handles:
      - Code blocks (```...```)
      - Bullet list lines (- item)
      - Numbered list lines (1. item)
      - Regular paragraphs
    """
    if not raw_answer:
        return

    # Split into code-block segments and regular text
    segments = re.split(r'(```[\w]*\n[\s\S]*?```)', raw_answer)

    for seg in segments:
        if seg.startswith("```"):
            # Code block
            code = re.sub(r'^```[\w]*\n', '', seg)
            code = re.sub(r'```$', '', code).strip()
            p = doc.add_paragraph(style="No Spacing")
            run = p.add_run(code)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1E, 0x37, 0x6D)
            p.paragraph_format.left_indent  = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
        else:
            # Regular text — split into lines
            lines = seg.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                # Bullet list
                if line.startswith("- ") or line.startswith("* "):
                    content = _clean_md_for_docx(line[2:])
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(content).font.size = Pt(10)
                # Numbered list
                elif re.match(r'^\d+\.\s', line):
                    content = _clean_md_for_docx(re.sub(r'^\d+\.\s', '', line))
                    p = doc.add_paragraph(style="List Number")
                    p.add_run(content).font.size = Pt(10)
                else:
                    content = _clean_md_for_docx(line)
                    if content:
                        p = doc.add_paragraph(style="No Spacing")
                        p.add_run(content).font.size = Pt(10)
                        p.paragraph_format.space_after = Pt(3)


def _add_page_break(doc: Document):
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# DATA LOADING
# ═══════════════════════════════════════════════════════════════════════════════

def load_questions() -> dict:
    """
    Returns dict:  topic → subtopic → list[Questions]
    """
    db = SessionLocal()
    try:
        rows = db.query(Questions).filter(
            Questions.role == "Data Scientist",
            Questions.verified == True,
        ).order_by(
            Questions.topic,
            Questions.subtopic,
            Questions.difficulty,
        ).all()
    finally:
        db.close()

    organized: dict = defaultdict(lambda: defaultdict(list))
    for row in rows:
        t = row.topic    or "General"
        s = row.subtopic or "General"
        organized[t][s].append(row)

    # Sort topics and subtopics
    return {
        t: dict(sorted(subs.items()))
        for t, subs in sorted(organized.items())
    }


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT GENERATION
# ═══════════════════════════════════════════════════════════════════════════════

def build_docx(organized: dict, output_path: str):
    doc = Document()

    # ── Page margins ───────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width   = Inches(8.5)
    section.page_height  = Inches(11)
    section.top_margin   = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin  = Inches(1.0)
    section.right_margin = Inches(1.0)

    # ── Normal style defaults ──────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ═══════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Data Science Interview")
    run.bold      = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x00, 0x3E, 0x8F)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle_p.add_run("Question & Answer Bank")
    run2.bold = True
    run2.font.size = Pt(28)
    run2.font.color.rgb = RGBColor(0x00, 0x3E, 0x8F)

    doc.add_paragraph()
    doc.add_paragraph()

    total_qs = sum(len(qs) for subs in organized.values() for qs in subs.values())
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run(
        f"500+ Curated Q&As across {len(organized)} topic areas\n"
        f"Sources: GitHub community repos + expert-curated additions\n"
        f"Generated: {datetime.now().strftime('%B %d, %Y')}"
    )
    info_run.font.size = Pt(13)
    info_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    doc.add_paragraph()

    # Topics overview table on cover
    topics_p = doc.add_paragraph()
    topics_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = topics_p.add_run("Topics Covered:")
    tr.bold = True
    tr.font.size = Pt(14)

    for topic, subs in organized.items():
        sub_count = sum(len(qs) for qs in subs.values())
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.add_run(f"• {topic}  ({sub_count} questions)").font.size = Pt(11)

    _add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════
    # SOURCES PAGE
    # ═══════════════════════════════════════════════════════════════
    sh = doc.add_heading("Data Sources", level=1)
    _set_heading_color(sh, RGBColor(0x00, 0x3E, 0x8F))

    sources = [
        ("1", "alexeygrigorev/data-science-interviews",
         "github.com/alexeygrigorev/data-science-interviews",
         "Community-maintained repo with theoretical ML/DL questions. "
         "Peer-reviewed; mistakes corrected via PRs. Covers 18+ ML topics."),
        ("2", "youssefHosni/Data-Science-Interview-Questions-Answers",
         "github.com/youssefHosni/Data-Science-Interview-Questions-Answers",
         "Curated Q&As covering ML, Deep Learning, Statistics, Probability, Python. "
         "Practitioner-authored and LinkedIn-community reviewed."),
        ("3", "iamtodor/data-science-interview-questions-and-answers",
         "github.com/iamtodor/data-science-interview-questions-and-answers",
         "Well-structured Q&A pairs on precision/recall, ROC, regularisation, "
         "feature selection, outlier handling, and more."),
        ("4", "kojino/120-Data-Science-Interview-Questions",
         "github.com/kojino/120-Data-Science-Interview-Questions",
         "Answers to 120 commonly asked DS interview questions, split across 7 categories: "
         "Probability, Statistical Inference, Data Analysis, Predictive Modeling, "
         "Programming, Product Metrics, Communication."),
        ("5", "roadmap.sh/questions/data-science",
         "roadmap.sh/questions/data-science",
         "Clean format covering beginner to advanced DS questions with practical examples."),
        ("6", "Expert-Curated Additions",
         "Internal / hand-authored",
         "Additional Q&As for SQL, Python best practices, Statistics, NLP, Deep Learning "
         "optimisation, MLOps, and Feature Engineering — crafted to fill topic gaps."),
    ]

    for num, name, url, desc in sources:
        p = doc.add_paragraph(style="No Spacing")
        p.paragraph_format.space_before = Pt(8)
        nr = p.add_run(f"[{num}] {name}\n")
        nr.bold = True
        nr.font.size = Pt(11)
        nr.font.color.rgb = RGBColor(0x00, 0x3E, 0x8F)
        ur = p.add_run(f"     {url}\n")
        ur.font.size = Pt(9)
        ur.font.color.rgb = RGBColor(0x22, 0x72, 0xB5)
        ur.italic = True
        dr = p.add_run(f"     {desc}")
        dr.font.size = Pt(10)

    _add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════════════════
    toc_h = doc.add_heading("Table of Contents", level=1)
    _set_heading_color(toc_h, RGBColor(0x00, 0x3E, 0x8F))

    for topic, subs in organized.items():
        tp = doc.add_paragraph(style="No Spacing")
        tp.paragraph_format.space_before = Pt(6)
        tr = tp.add_run(f"  {topic}")
        tr.bold = True
        tr.font.size = Pt(11)
        for subtopic, qs in subs.items():
            sp = doc.add_paragraph(style="No Spacing")
            sp.paragraph_format.left_indent = Inches(0.4)
            sp.add_run(f"    ▸ {subtopic}  ({len(qs)} questions)").font.size = Pt(10)

    _add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════
    # Q&A CONTENT  (Topic → Subtopic → Questions)
    # ═══════════════════════════════════════════════════════════════
    global_q_num = 1

    for topic, subs in organized.items():
        # Topic heading
        th = doc.add_heading(topic, level=1)
        _set_heading_color(th, RGBColor(0x00, 0x3E, 0x8F))

        for subtopic, qs in subs.items():
            # Subtopic heading
            sh2 = doc.add_heading(subtopic, level=2)
            _set_heading_color(sh2, RGBColor(0x1F, 0x5C, 0x99))

            for row in qs:
                diff   = (row.difficulty or "medium").lower()
                badge  = DIFF_BADGE.get(diff, "🔵 Medium")
                color  = DIFF_COLOR.get(diff, RGBColor(0x33, 0x66, 0x99))

                # Question header
                qh = doc.add_paragraph(style="No Spacing")
                qh.paragraph_format.space_before = Pt(10)
                qnum_run = qh.add_run(f"Q{global_q_num}. ")
                qnum_run.bold = True
                qnum_run.font.size = Pt(11)
                qnum_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x8F)
                qt_run = qh.add_run(row.question_text or "")
                qt_run.bold = True
                qt_run.font.size = Pt(11)

                # Difficulty badge
                badge_p = doc.add_paragraph(style="No Spacing")
                badge_p.paragraph_format.space_after = Pt(3)
                badge_run = badge_p.add_run(f"  {badge}")
                badge_run.font.size = Pt(9)
                badge_run.font.color.rgb = color
                badge_run.italic = True

                if row.question_type:
                    type_run = badge_p.add_run(f"  |  {row.question_type}")
                    type_run.font.size = Pt(9)
                    type_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    type_run.italic = True

                # Answer label
                ans_lbl = doc.add_paragraph(style="No Spacing")
                ans_lbl.paragraph_format.space_before = Pt(4)
                al_run = ans_lbl.add_run("Answer:")
                al_run.bold = True
                al_run.font.size = Pt(10)
                al_run.font.color.rgb = RGBColor(0x1A, 0x6E, 0x2A)

                # Answer body
                if row.ideal_answer:
                    _add_answer_block(doc, row.ideal_answer)
                else:
                    na_p = doc.add_paragraph(style="No Spacing")
                    na_p.add_run("No reference answer available yet.").font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

                # Keywords
                if row.keywords:
                    kw_p = doc.add_paragraph(style="No Spacing")
                    kw_p.paragraph_format.space_before = Pt(4)
                    kw_run_lbl = kw_p.add_run("Keywords: ")
                    kw_run_lbl.italic = True
                    kw_run_lbl.font.size = Pt(9)
                    kw_run_lbl.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                    kw_run = kw_p.add_run(row.keywords)
                    kw_run.italic = True
                    kw_run.font.size = Pt(9)
                    kw_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

                _add_horizontal_rule(doc)
                global_q_num += 1

        # Page break between major topics
        _add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════
    # APPENDIX — Difficulty Guide
    # ═══════════════════════════════════════════════════════════════
    ah = doc.add_heading("Appendix — Difficulty Guide", level=1)
    _set_heading_color(ah, RGBColor(0x00, 0x3E, 0x8F))

    levels = [
        ("🟢 Easy   (Fresher)",   "Conceptual definitions, common algorithms, basic metrics. "
                                   "Suitable for entry-level / fresher interviews."),
        ("🟡 Medium (Junior)",    "Application of concepts, comparisons, trade-offs, "
                                   "code reasoning. Suitable for 1–3 year experience."),
        ("🟠 Hard   (Mid-Level)", "Design decisions, advanced tuning, integration scenarios. "
                                   "Suitable for 3–5 year experience."),
        ("🔴 Expert (Senior)",    "Deep theoretical understanding, mathematical derivations, "
                                   "large-scale system design. Suitable for 5+ year senior roles."),
    ]
    for label, desc in levels:
        lp = doc.add_paragraph(style="No Spacing")
        lp.paragraph_format.space_before = Pt(8)
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(12)
        dp = doc.add_paragraph(style="No Spacing")
        dp.paragraph_format.left_indent = Inches(0.3)
        dp.add_run(desc).font.size = Pt(10)

    doc.save(output_path)
    print(f"\n✅  DOCX saved → {output_path}")
    print(f"   Total Q&As: {global_q_num - 1}")


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    print("\n" + "═" * 62)
    print("  Data Science Q&A Bank — DOCX Generator")
    print("═" * 62 + "\n")

    print("→ Loading questions from database …")
    organized = load_questions()

    total = sum(len(qs) for subs in organized.values() for qs in subs.values())
    print(f"  ✓ {total} questions across {len(organized)} topics\n")

    # Output alongside the scripts folder
    output_path = os.path.join(
        os.path.dirname(__file__),
        "..",
        "DS_Interview_QnA_Bank.docx",
    )
    output_path = os.path.abspath(output_path)
    # Fall back to a timestamped name if the file is locked (open in Word)
    if os.path.exists(output_path):
        try:
            open(output_path, "a").close()
        except PermissionError:
            from datetime import datetime
            ts = datetime.now().strftime("%H%M%S")
            output_path = output_path.replace(".docx", f"_{ts}.docx")

    print("→ Building DOCX …")
    build_docx(organized, output_path)


if __name__ == "__main__":
    main()
